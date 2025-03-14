"""
Описание класса Model. Model описывает основной функционал программы.

"""
import os
import re
import sys
import time
import base64
import openpyxl
from .view import View
from pathlib import Path
from docx import Document
from docx.shared import Pt
from datetime import datetime
from backend import database
from configparser import ConfigParser
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager



class Model():

    def __init__(self, version):
        self.db = None
        self.data = None
        self.data_update = []
        self.path_jpg = None
        self.path_ico = None
        self.path_excel = None
        self.path_config = None
        self.path_doc = None
        self.DRIVER = None
        self.SERVER_NAME = None
        self.DATABASE_NAME = None
        self.USERNAME = None
        self.PASSWORD = None
        self.driver_browser = None
        self.parse_config()
        View(self, version)

    def parse_config(self) -> None:
        self.path_jpg = self.get_resource_path("image.jpg")
        self.path_ico = self.get_resource_path("icon.ico")
        self.path_excel = self.get_resource_path("sample.xlsx")
        self.path_doc = self.get_resource_path("sample.docx")

    def get_resource_path(self, filename: str) -> str:
        path_progs = Path(__file__).parent
        """Возвращает корректный путь к файлу, работает и в `.exe`, и в обычном запуске Python"""
        if getattr(sys, '_MEIPASS', False):  # Проверяем, запущен ли скрипт в режиме PyInstaller
            return os.path.join(sys._MEIPASS, filename)
        path_file = os.path.join(path_progs, f"config\\{filename}")
        return path_file

    def connect_db(self) -> None:
        database.start_session(self.USERNAME, self.PASSWORD)

    def is_connect(self) -> bool:
        if database.check_connection():
            return True
        else:
            return False

    def find_in_db(self, surname, name, och, birthdate) -> None:
        self.data = database.fetch_person_by(surname=surname, name=name, patronymic=och, birthdate=birthdate)

    def is_data(self) -> bool:
        if not self.self.data:
            return True
        else:
            return False

    def find_in_db_for_check(self, surname, name, och, birthdate) -> None:
        self.data = database.get_person_for_check(surname, name, och, birthdate)

    def find_in_db_all_rows_for_check(self) -> None:
        self.data = database.get_all_rows_for_check()

    def completion_excel(self, file_out: str, row: int) -> None:


        # Загрузка книги
        wb = openpyxl.load_workbook(self.path_excel)
        # Работаем с активным листом (если нужен другой, уточните)
        sheet = wb.active

        # Начало заполнения таблицы
        # Фамилия на русском
        change_sheet(sheet, 10, 14, 4, self.data[row][0],  122)
        change_sheet(sheet, 176, 14, 4, self.data[row][0],  122)
        # Фамилия на латинице
        change_sheet(sheet, 12, 38, 4, self.data[row][1],  122)
        # Имя на русском
        change_sheet(sheet, 14, 14, 4, self.data[row][2],  122)
        change_sheet(sheet, 178, 14, 4, self.data[row][2],  122)
        # Имя на латинице
        change_sheet(sheet, 16, 34, 4, self.data[row][3],  122)
        # Отчество на русском
        change_sheet(sheet, 18, 30, 4, self.data[row][4],  122)
        change_sheet(sheet, 180, 34, 4, self.data[row][4],  122)
        # Отчество на латинице
        change_sheet(sheet, 20, 38, 4, self.data[row][5],  122)
        # Гражданство
        change_sheet(sheet, 24, 22, 4, self.data[row][6],  122)
        change_sheet(sheet, 182, 18, 4, self.data[row][6],  122)
        # Дата рождения
        if self.data[row][7] != '':
            date = self.data[row][7].strftime("%d.%m.%Y")
            # Число рождения
            change_sheet(sheet, 27, 30, 0, date[0],  30)
            change_sheet(sheet, 27, 34, 0, date[1],  34)
            change_sheet(sheet, 185, 30, 0, date[0],  30)
            change_sheet(sheet, 185, 34, 0, date[1],  34)
            # Месяц рождения
            change_sheet(sheet, 27, 46, 0, date[3],  46)
            change_sheet(sheet, 27, 50, 0, date[4],  50)
            change_sheet(sheet, 185, 50, 0, date[3],  50)
            change_sheet(sheet, 185, 54, 0, date[4],  54)
            # Год рождения
            change_sheet(sheet, 27, 58, 0, date[6],  58)
            change_sheet(sheet, 27, 62, 0, date[7],  62)
            change_sheet(sheet, 27, 66, 0, date[8],  66)
            change_sheet(sheet, 27, 70, 0, date[9],  70)
            change_sheet(sheet, 185, 66, 0, date[6],  66)
            change_sheet(sheet, 185, 70, 0, date[7],  70)
            change_sheet(sheet, 185, 74, 0, date[8],  74)
            change_sheet(sheet, 185, 78, 0, date[9],  78)
        # Пол
        if self.data[row][8] == "Мужской":
            change_sheet(sheet, 27, 90, 0, "X",  90)
            change_sheet(sheet, 185, 102, 0, "X",  102)
        elif self.data[row][8] == "Женский":
            change_sheet(sheet, 27, 106, 0, "X",  106)
            change_sheet(sheet, 185, 118, 0, "X",  118)
        # Государство рождения
        change_sheet(sheet, 29, 26, 4, self.data[row][9],  122)
        # Город рождения
        change_sheet(sheet, 31, 2, 4, self.data[row][10], 122)
        # Серия паспорта
        change_sheet(sheet, 37, 10, 4, self.data[row][11],  30)
        change_sheet(sheet, 187, 58, 4, self.data[row][11],  78)
        # Номер паспорта
        change_sheet(sheet, 37, 42, 4, self.data[row][12],  122)
        change_sheet(sheet, 187, 86, 4, self.data[row][12],  122)
        if self.data[row][13] != '':
            # Дата выдачи паспорта
            date = self.data[row][13].strftime("%d.%m.%Y")
            # Число
            change_sheet(sheet, 39, 9, 0, date[0],  9)
            change_sheet(sheet, 39, 13, 0, date[1],  13)
            change_sheet(sheet, 189, 9, 0, date[0],  9)
            change_sheet(sheet, 189, 13, 0, date[1],  13)
            # Месяц выдачи
            change_sheet(sheet, 39, 26, 0, date[3],  26)
            change_sheet(sheet, 39, 30, 0, date[4],  30)
            change_sheet(sheet, 189, 26, 0, date[3],  26)
            change_sheet(sheet, 189, 30, 0, date[4],  30)
            # Год выдачи
            change_sheet(sheet, 39, 38, 0, date[6],  38)
            change_sheet(sheet, 39, 42, 0, date[7],  42)
            change_sheet(sheet, 39, 46, 0, date[8],  46)
            change_sheet(sheet, 39, 50, 0, date[9],  50)
            change_sheet(sheet, 189, 38, 0, date[6],  38)
            change_sheet(sheet, 189, 42, 0, date[7],  42)
            change_sheet(sheet, 189, 46, 0, date[8],  46)
            change_sheet(sheet, 189, 50, 0, date[9],  50)
        if self.data[row][14] != '':
            # Дата срока действия паспорта
            date = self.data[row][14].strftime("%d.%m.%Y")
            # Число
            change_sheet(sheet, 39, 66, 0, date[0],  66)
            change_sheet(sheet, 39, 70, 0, date[1],  70)
            change_sheet(sheet, 189, 66, 0, date[0],  66)
            change_sheet(sheet, 189, 70, 0, date[1],  70)
            # Месяц
            change_sheet(sheet, 39, 82, 0, date[3],  82)
            change_sheet(sheet, 39, 86, 0, date[4],  86)
            change_sheet(sheet, 189, 82, 0, date[3],  82)
            change_sheet(sheet, 189, 86, 0, date[4],  86)
            # Год
            change_sheet(sheet, 39, 94, 0, date[6],  94)
            change_sheet(sheet, 39, 98, 0, date[7],  98)
            change_sheet(sheet, 39, 102, 0, date[8],  102)
            change_sheet(sheet, 39, 106, 0, date[9],  106)
            change_sheet(sheet, 189, 94, 0, date[6],  94)
            change_sheet(sheet, 189, 98, 0, date[7],  98)
            change_sheet(sheet, 189, 102, 0, date[8],  102)
            change_sheet(sheet, 189, 106, 0, date[9],  106)
        # Проверка на наличие визы, ВНЖ, РВПО
        if self.data[row][28] == "нет":
            # Наличие визы
            if self.data[row][15] == "1":
                change_sheet(sheet, 53, 8, 0, "X",  8)
            # Серия визы
            change_sheet(sheet, 56, 10, 4, self.data[row][16],  30)
            # Номер визы
            change_sheet(sheet, 56, 42, 4, self.data[row][17],  122)
            if self.data[row][18] != '':
                # Дата выдачи визы
                date = self.data[row][18].strftime("%d.%m.%Y")
                # Число
                change_sheet(sheet, 58, 9, 0, date[0],  9)
                change_sheet(sheet, 58, 13, 0, date[1],  13)
                # Месяц
                change_sheet(sheet, 58, 26, 0, date[3],  26)
                change_sheet(sheet, 58, 30, 0, date[4],  30)
                # Год
                change_sheet(sheet, 58, 38, 0, date[6],  38)
                change_sheet(sheet, 58, 42, 0, date[7],  42)
                change_sheet(sheet, 58, 46, 0, date[8],  46)
                change_sheet(sheet, 58, 50, 0, date[9],  50)
            if self.data[row][19] != '':
                # Дата срока визы
                date = self.data[row][19].strftime("%d.%m.%Y")
                # Число
                change_sheet(sheet, 58, 66, 0, date[0],  66)
                change_sheet(sheet, 58, 70, 0, date[1],  70)
                # Месяц
                change_sheet(sheet, 58, 82, 0, date[3],  82)
                change_sheet(sheet, 58, 86, 0, date[4],  86)
                # Год
                change_sheet(sheet, 58, 94, 0, date[6],  94)
                change_sheet(sheet, 58, 98, 0, date[7],  98)
                change_sheet(sheet, 58, 102, 0, date[8],  102)
                change_sheet(sheet, 58, 106, 0, date[9],  106)
        elif self.data[row][28] == "РВПО":
            change_sheet(sheet, 53, 108, 0, "X",  108)
            # Серия РВПО
            change_sheet(sheet, 56, 10, 4, self.data[row][31],  30)
            # Номер РВПО
            change_sheet(sheet, 56, 42, 4, self.data[row][32],  122)
            if self.data[row][29] != '':
                # Дата выдачи РВПО
                date = self.data[row][29]
                # Число
                change_sheet(sheet, 58, 9, 0, date[0],  9)
                change_sheet(sheet, 58, 13, 0, date[1],  13)
                # Месяц
                change_sheet(sheet, 58, 26, 0, date[3],  26)
                change_sheet(sheet, 58, 30, 0, date[4],  30)
                # Год
                change_sheet(sheet, 58, 38, 0, date[6],  38)
                change_sheet(sheet, 58, 42, 0, date[7],  42)
                change_sheet(sheet, 58, 46, 0, date[8],  46)
                change_sheet(sheet, 58, 50, 0, date[9],  50)

            if self.data[row][30] != '':
                # Дата срока РВПО
                date = self.data[row][30].strftime("%d.%m.%Y")
                # Число
                change_sheet(sheet, 58, 66, 0, date[0],  66)
                change_sheet(sheet, 58, 70, 0, date[1],  70)
                # Месяц
                change_sheet(sheet, 58, 82, 0, date[3],  82)
                change_sheet(sheet, 58, 86, 0, date[4],  86)
                # Год
                change_sheet(sheet, 58, 94, 0, date[6],  94)
                change_sheet(sheet, 58, 98, 0, date[7],  98)
                change_sheet(sheet, 58, 102, 0, date[8],  102)
                change_sheet(sheet, 58, 106, 0, date[9],  106)
        elif self.data[row][28] == "ВНЖ":
            change_sheet(sheet, 53, 36, 0, "X",  36)
            # Серия ВНЖ
            change_sheet(sheet, 56, 10, 4, self.data[row][31],  30)
            # Номер ВНЖ
            change_sheet(sheet, 56, 42, 4, self.data[row][32],  122)
            if self.data[row][29] != '':
                # Дата выдачи ВНЖ
                date = self.data[row][29]
                # Число
                change_sheet(sheet, 58, 9, 0, date[0],  9)
                change_sheet(sheet, 58, 13, 0, date[1],  13)
                # Месяц
                change_sheet(sheet, 58, 26, 0, date[3],  26)
                change_sheet(sheet, 58, 30, 0, date[4],  30)
                # Год
                change_sheet(sheet, 58, 38, 0, date[6],  38)
                change_sheet(sheet, 58, 42, 0, date[7],  42)
                change_sheet(sheet, 58, 46, 0, date[8],  46)
                change_sheet(sheet, 58, 50, 0, date[9],  50)
            if self.data[row][30] != '':
                # Дата срока ВНЖ
                date = self.data[row][30].strftime("%d.%m.%Y")
                # Число
                change_sheet(sheet, 58, 66, 0, date[0],  66)
                change_sheet(sheet, 58, 70, 0, date[1],  70)
                # Месяц
                change_sheet(sheet, 58, 82, 0, date[3],  82)
                change_sheet(sheet, 58, 86, 0, date[4],  86)
                # Год
                change_sheet(sheet, 58, 94, 0, date[6],  94)
                change_sheet(sheet, 58, 98, 0, date[7],  98)
                change_sheet(sheet, 58, 102, 0, date[8],  102)
                change_sheet(sheet, 58, 106, 0, date[9],  106)
        # Телефон
        # change_sheet(sheet, 62, 82, 4, self.data[row][20],  118)
        if self.data[row][21] != '':
            # Дата въезда
            date = self.data[row][21].strftime("%d.%m.%Y")
            # Число
            change_sheet(sheet, 66, 9, 0, date[0],  9)
            change_sheet(sheet, 66, 13, 0, date[1],  13)
            # Месяц
            change_sheet(sheet, 66, 26, 0, date[3],  26)
            change_sheet(sheet, 66, 30, 0, date[4],  30)
            # Год
            change_sheet(sheet, 66, 38, 0, date[6],  38)
            change_sheet(sheet, 66, 42, 0, date[7],  42)
            change_sheet(sheet, 66, 46, 0, date[8],  46)
            change_sheet(sheet, 66, 50, 0, date[9],  50)
        if self.data[row][22] != '':
            # Дата срок пребывания
            date = self.data[row][22].strftime("%d.%m.%Y")
            # Число
            change_sheet(sheet, 66, 66, 0, date[0],  66)
            change_sheet(sheet, 66, 70, 0, date[1],  70)
            change_sheet(sheet, 222, 49, 0, date[0],  49)
            change_sheet(sheet, 222, 53, 0, date[1],  53)
            change_sheet(sheet, 266, 38, 0, date[0],  38)
            change_sheet(sheet, 266, 42, 0, date[1],  42)
            # Месяц
            change_sheet(sheet, 66, 82, 0, date[3],  82)
            change_sheet(sheet, 66, 86, 0, date[4],  86)
            change_sheet(sheet, 222, 65, 0, date[3],  65)
            change_sheet(sheet, 222, 69, 0, date[4],  69)
            change_sheet(sheet, 266, 54, 0, date[3],  54)
            change_sheet(sheet, 266, 58, 0, date[4],  58)
            # Год
            change_sheet(sheet, 66, 94, 0, date[6],  94)
            change_sheet(sheet, 66, 98, 0, date[7],  98)
            change_sheet(sheet, 66, 102, 0, date[8],  102)
            change_sheet(sheet, 66, 106, 0, date[9],  106)
            change_sheet(sheet, 222, 79, 0, date[6],  79)
            change_sheet(sheet, 222, 83, 0, date[7],  83)
            change_sheet(sheet, 222, 87, 0, date[8],  87)
            change_sheet(sheet, 222, 91, 0, date[9],  91)
            change_sheet(sheet, 266, 68, 0, date[6],  68)
            change_sheet(sheet, 266, 72, 0, date[7],  72)
            change_sheet(sheet, 266, 76, 0, date[8],  76)
            change_sheet(sheet, 266, 80, 0, date[9],  80)
        # Миграционная карта серия
        change_sheet(sheet, 68, 34, 4, self.data[row][23],  46)
        # Миграционная карта номер
        change_sheet(sheet, 68, 54, 4, self.data[row][24],  102)

        # Номер корпус
        if self.data[row][25]:
            num = self.data[row][25]
            sheet.cell(row=104, column=46, value=f"КОРПУС {num[0]}")
            sheet.cell(row=200, column=46, value=f"КОРПУС {num[0]}")

        if self.data[row][26] != '':
            # Дата начала найма
            date = self.data[row][26]
            # Число найма
            change_sheet(sheet, 112, 90, 0, date[0],  90)
            change_sheet(sheet, 112, 94, 0, date[1],  94)
            # Месяц найма
            change_sheet(sheet, 112, 102, 0, date[3],  102)
            change_sheet(sheet, 112, 106, 0, date[4],  106)
            # Год найма
            change_sheet(sheet, 112, 114, 0, date[8],  114)
            change_sheet(sheet, 112, 118, 0, date[9],  118)
            # Номер договора
            change_sheet(sheet, 114, 46, 4, self.data[row][27],  90)
        if self.data[row][42]:
            pattern_reg = r"РЕСПУБЛИКА\s[А-ЯЁ]+(?:\s[А-ЯЁ]+)?|[А-ЯЁ]+(?:-[А-ЯЁ]+)?\s(?:ОБЛАСТЬ|КРАЙ|РЕСПУБЛИКА|АО)"
            pattern_rayon = r"[А-ЯЁ]+(?:-[А-ЯЁ]+)?\s(?:РАЙОН)"
            pattern_city = r"(?:Г|ПГТ|РП|КП|К|ДП|П|НП|С|М|Д|СЛ|СТ|СТ-ЦА|Х|КЛХ|СВХ)\.\s?\b[А-ЯЁ]+(?:-[А-ЯЁ]+)?\b"
            pattern_street = r"(?:УЛ|АЛЛЕЯ|БУЛЬВАР|МАГИСТРАЛЬ|ПЕРЕУЛОК|ПЛОЩАДЬ|ПРОЕЗД|ПРОСПЕКТ|ПРОУЛОК|РАЗЪЕЗД|СПУСК|ТРАКТ|ТУПИК|ШОССЕ|УЛИЦА)\.?\s?[А-ЯЁ]+(?:-[А-ЯЁ]+)?|[А-ЯЁ]+(?:-[А-ЯЁ]+)?\s(?:УЛ|ПР|П|ПРОЕЗД|УЛИЦА|ПРОСПЕКТ|БУЛЬВАР)\.?"
            pattern_dom = r"Д.\s?\b[0-9]+(?:[А-ЯЁ]+)?\b"
            pattern_korp = r"КОРП.\s?\b[0-9]+\b|КОРП.\s?\b[А-ЯЁ]+\b"
            pattern_str = r"СТР.\s?\b[0-9]+\b|СТР.\s?\b[А-ЯЁ]+\b"
            pattern_kv = r"КВ.\s?\b[0-9]+\b"
            reg = "".join(re.findall(pattern_reg, self.data[row][42]))
            change_sheet(sheet, 83, 2, 4, reg, 122)
            rayon = "".join(re.findall(pattern_rayon, self.data[row][42]))
            change_sheet(sheet, 85, 2, 4, rayon, 122)
            city = "".join(re.findall(pattern_city, self.data[row][42]))
            change_sheet(sheet, 87, 2, 4, city, 122)
            street = "".join(re.findall(pattern_street, self.data[row][42]))
            change_sheet(sheet, 89, 2, 4, street, 122)
            dom = "".join(re.findall(pattern_dom, self.data[row][42]))
            sheet.cell(row=91, column=2, value=f"ДОМ {dom.replace('Д. ','')}")
            korp = "".join(re.findall(pattern_korp, self.data[row][42]))
            str_ = "".join(re.findall(pattern_str, self.data[row][42]))
            if korp:
                sheet.cell(row=91, column=46,
                        value=f"КОРПУС {korp.replace('КОРП. ','')}")
            elif str_:
                sheet.cell(row=91, column=46,
                        value=f"СТРОЕНИЕ {str_.replace('СТР. ','')}")
            kv = "".join(re.findall(pattern_kv, self.data[row][42]))
            if kv:

                sheet.cell(row=93, column=2,
                        value=f"КВАРТИРА {kv.replace('КВ. ','')}")

        # Сохранение изменений
        if os.path.isdir(".\\out"):
            pass
        else:
            os.mkdir(f".\\out")

        output_path = f".\\out\\{file_out}.xlsx"
        wb.save(output_path)
        os.startfile(output_path)

    def complection_word(self, file_out: str, row: int, check: int) -> None:

        doc = Document(self.path_doc)
        if check == 1:
            UpdateWord(doc, 1, 12, 1, "X")
        elif check == 2:
            UpdateWord(doc, 1, 14, 1, "X")
        elif check == 3:
            UpdateWord(doc, 1, 16, 1, "X")
        UpdateWord(doc, 1, 19, 3, self.data[row][0])  # Фамилия на русском
        UpdateWord(doc, 1, 19, 15, self.data[row][1])  # Фамилия на латинице
        UpdateWord(doc, 1, 21, 4, self.data[row][2])  # Имя на русском
        UpdateWord(doc, 1, 21, 15, self.data[row][3])  # Имя на латинице
        UpdateWord(doc, 1, 23, 7, self.data[row][4])  # Отчество на русском
        UpdateWord(doc, 1, 23, 15, self.data[row][5])  # Отчество на латинице
        if self.data[row][7] != '':
            UpdateWord(doc, 1, 25, 5, self.data[row][7].strftime("%d.%m.%Y"))  # Дата рождения
        # Стаана и город рождения
        UpdateWord(doc, 1, 25, 13,
                f"{self.data[row][9]}, {self.data[row][10]}", 11)
        UpdateWord(doc, 1, 27, 8, self.data[row][6])  # Гражданство
        if self.data[row][8] == "Мужской":
            UpdateWord(doc, 1, 27, 20, "X")  # Пол
        elif self.data[row][8] == "Женский":
            UpdateWord(doc, 1, 27, 22, "X")  # Пол
        UpdateWord(doc, 1, 29, 2, self.data[row][11])  # Серия паспорта
        UpdateWord(doc, 1, 29, 6, self.data[row][12])  # Номер паспорта
        if self.data[row][13] != '':
            UpdateWord(doc, 1, 29, 15, self.data[row][13].strftime("%d.%m.%Y"), 11)  # Дата выдачи паспорта
        if self.data[row][14] != '':
            # Дата срока действия паспорта
            UpdateWord(doc, 1, 29, 19, self.data[row][14].strftime("%d.%m.%Y"), 11)
        if self.data[row][25]:
            num = self.data[row][25]
            UpdateWord(
                doc, 2, 2, 0, f"Г. МОСКВА, КОЧНОВСКИЙ ПР., Д.7, КОРПУС {num[0]}")
        if self.data[row][21] != '':
            UpdateWord(doc, 2, 3, 14, self.data[row][21].strftime("%d.%m.%Y"))  # Дата въезда
        if self.data[row][22] != '':
            UpdateWord(doc, 2, 3, 23, self.data[row][22].strftime("%d.%m.%Y"))  # Срок пребывания
        if self.data[row][28] == "нет":
            UpdateWord(doc, 2, 6, 4, self.data[row][33], 12, False)  # Кратност визы
            UpdateWord(doc, 2, 6, 28, "учебная", 12, False)  # Категория визы
            UpdateWord(doc, 2, 7, 3, "учеба", 12, False)  # Цель визы
            UpdateWord(doc, 2, 7, 19, self.data[row][16])  # Серия визы
            UpdateWord(doc, 2, 7, 28, self.data[row][17])  # Номер визы
            # Инденитификационный номер визы
            UpdateWord(doc, 2, 8, 8, self.data[row][34])
            if self.data[row][18] != '':
                UpdateWord(doc, 2, 8, 24, self.data[row][18].strftime("%d.%m.%Y"))  # Дата выдачи визы
            if self.data[row][19] != '':
                UpdateWord(doc, 2, 8, 31, self.data[row][19].strftime("%d.%m.%Y"))  # Дата срока визы
        elif self.data[row][28] == "РВПО":
            if self.data[row][16] != "" or self.data[row][16].isspace():
                UpdateWord(doc, 2, 6, 4, self.data[row][33], 12, False)  # Кратност визы
                UpdateWord(doc, 2, 6, 28, "учебная", 12, False)  # Категория визы
                UpdateWord(doc, 2, 7, 3, "учеба", 12, False)  # Цель визы
                UpdateWord(doc, 2, 7, 19, self.data[row][16])  # Серия визы
                UpdateWord(doc, 2, 7, 28, self.data[row][17])  # Номер визы
                # Инденитификационный номер визы
                UpdateWord(doc, 2, 8, 8, self.data[row][34])
                if self.data[row][18] != '':
                    UpdateWord(doc, 2, 8, 24, self.data[row][18].strftime("%d.%m.%Y"))  # Дата выдачи визы
                if self.data[row][19] != '':
                    UpdateWord(doc, 2, 8, 31, self.data[row][19].strftime("%d.%m.%Y"))  # Дата срока визы
        elif self.data[row][28] == "ВНЖ":
            if self.data[row][16] != "" or self.data[row][16].isspace():
                UpdateWord(doc, 2, 6, 4, self.data[row]
                        [33], 12, False)  # Кратност визы
                UpdateWord(doc, 2, 6, 28, "учебная", 12, False)  # Категория визы
                UpdateWord(doc, 2, 7, 3, "учеба", 12, False)  # Цель визы
                UpdateWord(doc, 2, 7, 19, self.data[row][16])  # Серия визы
                UpdateWord(doc, 2, 7, 28, self.data[row][17])  # Номер визы
                # Инденитификационный номер визы
                UpdateWord(doc, 2, 8, 8, self.data[row][34])
                if self.data[row][18] != '':
                    UpdateWord(doc, 2, 8, 24, self.data[row][18].strftime("%d.%m.%Y"))  # Дата выдачи визы
                if self.data[row][19] != '':
                    UpdateWord(doc, 2, 8, 31, self.data[row][19].strftime("%d.%m.%Y"))  # Дата срока визы
        if self.data[row][38] == "":
            year_kon_st = 0
        else:
            date = self.data[row][38].strftime("%d.%m.%Y")
            year_kon_st = int(date[6]+date[7]+date[8]+date[9])
        if self.data[row][40] == "":
            year_gos_st = 0
        else:
            date = self.data[row][40].strftime("%d.%m.%Y")
            year_gos_st = int(date[6]+date[7]+date[8]+date[9])
        if year_gos_st > year_kon_st:
            UpdateWord(doc, 2, 12, 10, "гос.направление", 11, False)  # Направление
            if self.data[row][40] != '':
                # Дата выдачи контракта
                UpdateWord(doc, 2, 12, 21, self.data[row][40].strftime("%d.%m.%Y"), 11)
            UpdateWord(doc, 2, 12, 28, self.data[row][35], 11)  # Номер контракта
            if self.data[row][40] != '':
                UpdateWord(doc, 2, 14, 5, self.data[row][40].strftime("%d.%m.%Y"), 11)  # Срок обучения с
            if self.data[row][41] != '':
                UpdateWord(doc, 2, 14, 15, self.data[row][41].strftime("%d.%m.%Y"), 11)  # Срок обучения по
        else:
            UpdateWord(doc, 2, 12, 10, "контракт", 11, False)  # Направление
            if self.data[row][38] != '':
                # Дата выдачи контракта
                UpdateWord(doc, 2, 12, 21, self.data[row][38].strftime("%d.%m.%Y"), 11)
            UpdateWord(doc, 2, 12, 28, self.data[row][37], 11)  # Номер контракта
            if self.data[row][38] != '':
                UpdateWord(doc, 2, 14, 5, self.data[row][38].strftime("%d.%m.%Y"), 11)  # Срок обучения с
            if self.data[row][39] != '':
                UpdateWord(doc, 2, 14, 15, self.data[row][39].strftime("%d.%m.%Y"), 11)  # Срок обучения по

        # Сохранение изменений
        if os.path.isdir(".\\out"):
            pass
        else:
            os.mkdir(f".\\out")

        output_path = f".\\out\\{file_out}.docx"
        doc.save(output_path)
        os.startfile(output_path)

    def check_on_gosuslugi(self, row: int) -> None:
        if self.check_open_browser() == False:
            self.open_browser()
        # Задержка для загрузки страницы
        time.sleep(3)

        while True:
            try:
                input_birth_date = self.driver_browser.find_element(By.NAME, "c_birth_date")  # Поле даты рождения
                if input_birth_date:
                    input_birth_date.clear()  # Очистить поле
                    input_birth_date.send_keys(self.data[row][4])  # Вставить дату (пример: 01.01.1990)
                    break
            except Exception as e:
                print("Ошибка:", e)
                if "10054" in str(e) or self.check_open_browser() == False:
                    return


        while True:
            try:
                input_pass_ser = self.driver_browser.find_element(By.NAME, "c_series")
                if input_pass_ser:
                    input_pass_ser.clear()  # Очистить поле
                    input_pass_ser.send_keys(self.data[row][5])
                    input_pass_num = self.driver_browser.find_element(By.NAME, "c_number")
                    input_pass_num.clear()  # Очистить поле
                    input_pass_num.send_keys(self.data[row][6])
                    input_pass_date = self.driver_browser.find_element(By.NAME, "c_issue_date")
                    input_pass_date.clear()  # Очистить поле
                    input_pass_date.send_keys(self.data[row][7])
                    break
            except Exception as e:
                print("Ошибка:", e)
                if "10054" in str(e) or self.check_open_browser() == False:
                    return

        while True:
            try:
                is_class = self.driver_browser.find_element(By.CLASS_NAME, "info__img")
                if is_class:
                    self.save_as_pdf(row)
                    break
            except Exception as e:
                print("Ошибка:", e)
                if "10054" in str(e) or self.check_open_browser() == False:
                    return

    def open_browser(self):
        try:
            # Настройка WebDriver
            options = webdriver.ChromeOptions()
            options.add_experimental_option("detach", True)  # Оставить браузер открытым
            self.driver_browser = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)

            # Открываем страницу (замените URL на нужный)
            self.driver_browser.get("https://www.gosuslugi.ru/655781/1/form")
        except Exception as e:
                print("Ошибка:", e)
                raise RuntimeError(f"Произошла ошибка ! {e}")  # Исключение с сообщением

    def close_browser(self):
        # Оставляем браузер открытым на 10 секунд
        time.sleep(10)
        self.driver_browser.quit()

    def check_open_browser(self):
        try:
            if self.driver_browser is None:
                return False
            if self.driver_browser.service.process.poll() is not None:
                return False
            if len(self.driver_browser.window_handles) == 0:
                return False
            return True
        except:
            return False

    def save_as_pdf(self, row):
        if os.path.isdir(".\\out_check"):
            pass
        else:
            os.mkdir(f".\\out_check")

        output_path = f".\\out_check\\{self.data[row][0]}.pdf"
        """Сохраняет страницу в PDF через Chrome DevTools Protocol"""
        pdf = self.driver_browser.execute_cdp_cmd("Page.printToPDF", {})
        with open(output_path, "wb") as f:
            f.write(base64.b64decode(pdf['data']))

    def process_update_database(self, window, row):
        self.copy_change_to_data_update(window, row)
        database.update_person(self.data[row][60], self)

    def copy_change_to_data_update(self, window, row):
        self.data_update.clear()
        # print(self.data)
        # print(self.data_update)
        self.append_change(window.fru.get(), row, 0)
        self.append_change(window.last_lat.get(), row, 1)
        self.append_change(window.name_rus.get(), row, 2)
        self.append_change(window.nla.get(), row, 3)
        self.append_change(window.och.get(), row, 4)
        self.append_change(window.oche.get(), row, 5)
        self.append_change(window.ctz1.get(), row, 6)
        self.append_change(window.dob.get(), row, 7)
        self.append_change(window.sex.get(), row, 8)
        self.append_change(window.pob.get(), row, 9)
        self.append_change(window.cob.get(), row, 10)
        self.append_change(window.pas_ser.get(), row, 11)
        self.append_change(window.pas_num.get(), row, 12)
        self.append_change(window.pds.get(), row, 13)
        self.append_change(window.pde.get(), row, 14)
        self.append_change('', row, 15)
        self.append_change(window.vis_ser.get(), row, 16)
        self.append_change(window.vis_num.get(), row, 17)
        self.append_change(window.vis_start.get(), row, 18)
        self.append_change(window.vis_end.get(), row, 19)
        self.append_change(window.tel_nom.get(), row, 20)
        self.append_change(window.d_enter.get(), row, 21)
        self.append_change('', row, 22)
        self.append_change(window.mcs.get(), row, 23)
        self.append_change(window.mcn.get(), row, 24)
        self.append_change(window.k.get(), row, 25)
        self.append_change(window.dnd.get(), row, 26)
        self.append_change(window.dog_obsh.get(), row, 27)
        self.append_change(window.rf.get(), row, 28)
        self.append_change(window.rfd.get(), row, 29)
        self.append_change(window.mot.get(), row, 30)
        self.append_change(window.ser.get(), row, 31)
        self.append_change(window.nmr.get(), row, 32)
        self.append_change(window.vis_krat.get(), row, 33)
        self.append_change(window.vis_id.get(), row, 34)
        self.append_change(window.gos_nap.get(), row, 35)
        self.append_change(window.proz.get(), row, 36)
        self.append_change(window.kontrakt.get(), row, 37)
        self.append_change(window.kont_start.get(), row, 38)
        self.append_change(window.kont_end.get(), row, 39)
        self.append_change('', row, 40)
        self.append_change('', row, 41)
        self.append_change(window.star.get(), row, 42)
        self.append_change(window.d_poluch.get(), row, 43)
        self.append_change(window.str_poluch.get(), row, 44)
        self.append_change(window.city_poluch.get(), row, 45)
        self.append_change(window.num_prig.get(), row, 46)
        self.append_change(window.kpp.get(), row, 47)
        self.append_change('', row, 48)
        self.append_change(window.mot.get(), row, 49)
        self.append_change(window.uch_st_st.get(), row, 50)
        self.append_change(window.pr.get(), row, 51)
        self.append_change(window.fr.get(), row, 52)
        self.append_change(window.o_p.get(), row, 53)
        self.append_change(window.prikaz.get(), row, 54)
        self.append_change(window.prik_start.get(), row, 55)
        self.append_change(window.o_s.get(), row, 56)
        self.append_change(window.d_naym.get(), row, 57)
        self.append_change(window.email.get(), row, 58)
        self.append_change(window.prim.get(), row, 59)
        self.append_change('', row, 60)
        self.check_update_database(row)

    def append_change(self, value, row, col):
        if value == '':
            self.data_update.append(self.data[row][col])
        else:
            if isinstance(value, datetime):
                value = datetime.strptime(self.data_update[row][col], "%d.%m.%Y")
            self.data_update.append(value)

    def check_update_database(self, row):
        for i, var in enumerate(self.data_update):
            if var == self.data[row][i]:
                self.data_update[i] = ''
        print(self.data_update)

def change_sheet(
                sheet,
                row: int,
                col: int,
                step: int,
                value: str,
                max_col: int
) -> None:
    col_cur = col
    row_cur = row
    for val in value:
        if col_cur <= max_col:
            sheet.cell(row=row_cur, column=col_cur, value=val)
            col_cur += step

def UpdateWord(
                    doc,
                    table_index: int,
                    row_index: int,
                    cell_index: int,
                    new_text: str,
                    font_size = 12 ,
                    uppercase = True
        ) -> None:
            """ Обновляет текст в ячейке таблицы, добавляя жирность и изменение регистра """
            try:
                table = doc.tables[table_index]
                if row_index < len(table.rows) and cell_index < len(table.rows[row_index].cells):
                    cell = table.rows[row_index].cells[cell_index]

                    # Определяем окончательный текст
                    formatted_text = new_text.upper() if uppercase else new_text.lower()

                    # Если в ячейке нет параграфов, создаем новый
                    if not cell.paragraphs:
                        paragraph = cell.add_paragraph()
                    else:
                        paragraph = cell.paragraphs[0]

                    # Если нет runs, создаем новый run
                    if not paragraph.runs:
                        run = paragraph.add_run(formatted_text)
                    else:
                        run = paragraph.runs[0]
                        run.text = formatted_text

                    # Применяем жирный стиль
                    run.bold = True
                    run.font.name = "Times New Roman"  # Устанавливаем шрифт
                    run.font.size = Pt(font_size)  # Устанавливаем размер шрифта
                else:
                    print(f"Invalid row or cell index: ({row_index}, {cell_index})")
            except IndexError as e:
                print(f"Error accessing table: {e}")
