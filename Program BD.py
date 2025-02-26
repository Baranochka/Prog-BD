"""
Программа для работы с базой данных "Название" и последующего вывода в специальной форме в Excel
P.S. Извините за кринж
"""
__version__ = "v0.7"  # Надо менять версию после каждого изменения

import os
import sys
import re
import openpyxl
import customtkinter as ctk
from PIL import Image
from tkinter import ttk
from pathlib import Path
from docx import Document
from backend import MSSQL
from docx.shared import Pt
from datetime import datetime
from backend import UpdateApp
from configparser import ConfigParser


data = []
# Создание объекта для работы с базой данных
db = None
# Переменная для хранения строки таблицы, на которую нажали
row_click = 0

def get_resource_path(filename: str) -> str:
    path_progs = Path(__file__).parent
    """Возвращает корректный путь к файлу, работает и в `.exe`, и в обычном запуске Python"""
    if getattr(sys, '_MEIPASS', False):  # Проверяем, запущен ли скрипт в режиме PyInstaller
        return os.path.join(sys._MEIPASS, filename)
    path_file = os.path.join(path_progs, f"config\\{filename}")
    return path_file


jpg_path = get_resource_path("image.jpg")
ico_path = get_resource_path("icon.ico")
excle_path = get_resource_path("sample.xlsx")
config_path = get_resource_path("config.ini")
doc_path = get_resource_path("sample.docx")

config = ConfigParser()
config.read(config_path)

# DRIVER = config["SQL Server"]["DRIVER"]
# SERVER_NAME = config["SQL Server"]["SERVER_NAME"]
# DATABASE_NAME = config["SQL Server"]["DATABASE_NAME"]
# USERNAME = None
# PASSWORD = None


# DRIVER = config["Test"]["DRIVER"]
# SERVER_NAME = config["Test"]["SERVER_NAME"]
# DATABASE_NAME = config["Test"]["DATABASE_NAME"]
# USERNAME      = None
# PASSWORD      = None


DRIVER = config["Test_Maks"]["DRIVER"]
SERVER_NAME = config["Test_Maks"]["SERVER_NAME"]
DATABASE_NAME = config["Test_Maks"]["DATABASE_NAME"]
USERNAME = None
PASSWORD = None


class WindowAuthorizationConnectionBD(ctk.CTkToplevel):
    def __init__(self, parent: ctk.CTk):
        super().__init__(parent)

        self.parent = parent
        # self.attributes("-topmost", True)
        self.title("Авторизация")
        center_window(self, 300, 200)
        self.configure(fg_color="white")
        self.iconbitmap(ico_path)

        self.frame_auth = ctk.CTkFrame(self, fg_color="white")
        self.frame_auth.pack(expand=True, fill=ctk.BOTH)
        ctk.CTkLabel(self.frame_auth, text="Авторизация",
                      bg_color="white", font=("Arial", 17)).place(x=100, y=10)
        ctk.CTkLabel(self.frame_auth, text="Логин",
                      bg_color="white", height=15).place(x=130, y=40)
        self.entry_login = ctk.CTkEntry(
            self.frame_auth, width=180, justify="center")

        self.entry_login.place(x=60, y=60)
        ctk.CTkLabel(self.frame_auth, text="Пароль",
                      bg_color="white", height=10).place(x=126, y=90)
        self.entry_password = ctk.CTkEntry(
            self.frame_auth, width=180, justify="center")

        self.entry_password.place(x=60, y=110)
        self.entry_password.configure(show="*")

        self.frame_incorrect = ctk.CTkFrame(self, fg_color="white")

        ctk.CTkLabel(self.frame_incorrect, text="Соединение не установлено",
                      fg_color="white", text_color="red").place(x=75, y=75)
        self.frame_incorrect.pack_forget()

        self.frame_correct = ctk.CTkFrame(self, fg_color="white")

        ctk.CTkLabel(self.frame_correct, text="Соединение установлено",
                      fg_color="white", text_color="green").place(x=75, y=75)
        self.frame_correct.pack_forget()

        self.entry_login.bind(
            "<FocusIn>", lambda e: self.entry_login.delete(0, ctk.END))
        self.entry_password.bind(
            "<FocusIn>", lambda e: self.entry_password.delete(0, ctk.END))

        ctk.CTkButton(self.frame_auth, text="Подключиться",
                       command=self.connect).place(x=80, y=150)
        # Перехватываем нажатие на крестик (закрытие окна)
        self.protocol("WM_DELETE_WINDOW", self.close_app)

    def close_app(self):
        # Завершает весь Tkinter
        self.parent.quit()

    def connect(self):
        self.frame_auth.pack_forget()
        global DRIVER, SERVER_NAME, DATABASE_NAME, USERNAME, PASSWORD
        USERNAME = self.entry_login.get()
        PASSWORD = self.entry_password.get()
        global db
        db = MSSQL(DRIVER, SERVER_NAME, DATABASE_NAME, USERNAME, PASSWORD)
        if db.is_connected == True:
            self.frame_correct.pack(expand=True, fill=ctk.BOTH)
            self.after(1000, self.hide_correct_frame)

        else:
            self.frame_incorrect.pack(expand=True, fill=ctk.BOTH)
            self.after(1000, self.hide_incorrect_frame)

    def hide_incorrect_frame(self):
        self.frame_incorrect.pack_forget()
        self.frame_auth.pack(expand=True, fill=ctk.BOTH)
        self.entry_login.delete(0, ctk.END)
        self.entry_password.delete(0, ctk.END)

    def hide_correct_frame(self):

        self.destroy()
        self.parent.deiconify()


class WindowProgramm(ctk.CTk):
    # Инициализация окна программы и откликов в программе
    def __init__(window):
        super().__init__()
        # Описание окна
        window.title("Программа для работы с БД")
        center_window(window, 1000, 600)
        window.configure(fg_color="white")
        # Устанавливаем иконку (Только .ico!)
        window.iconbitmap(ico_path)

        # Создаём CTkLabel с изображением
        my_image = ctk.CTkImage(
            light_image=Image.open(jpg_path), size=(150, 100))

        ctk.CTkLabel(window, image=my_image, text="").place(x=10, y=10)

        """--------------------------------------------------ФИО---------------------------------------------------"""

        # Создание ярлыка и текстового поля для ввода фамилии
        ctk.CTkLabel(master=window, text="Фамилия:",
                      bg_color="white").place(x=20, y=150)
        window.entry_familia = ctk.CTkEntry(master=window, width=250)
        window.entry_familia.place(x=90, y=150)

        # # Создание ярлыка и текстового поля для ввода имени
        ctk.CTkLabel(master=window, text="Имя:",
                      bg_color="white").place(x=20, y=180)
        window.entry_name = ctk.CTkEntry(master=window, width=250)
        window.entry_name.place(x=90, y=180)

        # # Создание ярлыка и текстового поля для ввода отчества
        ctk.CTkLabel(master=window, text="Отчество:",
                      bg_color="white").place(x=20, y=210)
        window.entry_otchestvo = ctk.CTkEntry(master=window, width=250)
        window.entry_otchestvo.place(x=90, y=210)

        """---------------------------------------------------ДАТА--------------------------------------------------"""

        # Создание ярлыка даты рождения
        ctk.CTkLabel(window, text="Дата рождения:",
                      bg_color="white").place(x=20, y=240)

        # Переменные для хранения значений даты рождения
        window.day_var = ctk.StringVar(value="1")
        window.month_var = ctk.StringVar(value="1")
        window.year_var = ctk.StringVar(value="1900")

        # CTkSpinbox для дня от 1 до 31
        ctk.CTkEntry(master=window, textvariable=window.day_var,
                      width=20).place(x=130, y=240)

        # CTkSpinbox для месяца от 1 до 12
        ctk.CTkEntry(master=window, textvariable=window.month_var,
                      width=20).place(x=160, y=240)

        # CTkSpinbox для года от 1900 до 2025
        ctk.CTkEntry(master=window, textvariable=window.year_var,
                      width=45).place(x=190, y=240)

        """---------------------------------------------------КНОПКА-ПОИСКА--------------------------------------------------------------"""

        # Создание кнопки "Поиск"
        ctk.CTkButton(window, text="Поиск",
                       command=window.click_find).place(x=450, y=240)

        """------------------------------------------------------ТАБЛИЦА-------------------------------------------------------------"""

        # Создание контейнера для таблицы (изначально скрыт)
        window.frame_table = ctk.CTkFrame(
            window, fg_color="white", width=1000, height=310)

        # Отключение видимости таблицы
        window.frame_table.pack_forget()

        # Первая строка с названиями столбцов
        window.columns = ("Фамилия",  # 0      fru
                          "Familia",  # 1      last_lat
                          "Имя",  # 2      name_rus
                          "Imya"  # 3      nla
                          #   "Отчество",                               #4      och
                          #   "Otchestvo",                              #5      oche
                          #   "Гражданство",                            #6      pob
                          #   "Дата рождения",                          #7      dob
                          #   "Пол",                                    #8      sex
                          #   "Государство рождения",                   #9      pob
                          #   "Город рождения",                         #10     cob
                          #   "Серия паспорта",                         #11     pas_ser
                          #   "Номер паспорта",                         #12     pas_num
                          #   "Дата выдачи паспорта",                   #13     pds
                          #   "Срок действия паспорта",                 #14     pde
                          #   "Признак наличия визы",                   #15     visa_priz
                          #   "Виза серия",                             #16     vis_ser
                          #   "Виза номер",                             #17     vis_num
                          #   "Дата выдачи визы",                       #18     d_poluch
                          #   "Срок действия визы",                     #19     vis_end
                          #   "Телефон",                                #20     tel_nom
                          #   "Дата въезда",                            #21     d_enter
                          #   "Срок пребывания до",                     #22     date_okon
                          #   "Миграционная карта серия",               #23     mcs
                          #   "Миграционная карта номер",               #24     mcn
                          #   "Номер корпуса общежития",                #25     obsch
                          #   "Дата начала договора с общагой",         #26     dnd
                          #   "Номер договора",                         #27     dog_obsh
                          #   "Наличие ВНЖ, РВПО",                      #28     rf
                          #   "Дата выдачи ВНЖ, РВПО",                  #29     rfd
                          #   "Срок действия ВНЖ, РВПО",                #30     mot
                          #   "Серия ВНЖ, РВПО",                        #31     ser
                          #   "Номер ВНЖ, РВПО",                        #32     nmr
                          #    "Кратность визы",                        #33     vis_krat
                          #    "Индентификатор визы",                   #34     vis_id
                          #    "Направление",                           #35     gos_nap
                          #    "Дата выдачи контракта",                 #36     gos_start
                          #    "Номер контракта",                       #37     kontrakt
                          #    "Срок обучения с",                       #38     kont_start
                          #    "Срок обучения по",                      #39     kont_end
                          )
        style = ttk.Style()
        style.configure("Custom.Treeview", font=("Arial", 14))
        # Создание таблицы
        window.tree = ttk.Treeview(window.frame_table, columns=window.columns,
                                   height=20, show='headings', style="Custom.Treeview")

        # Определение размера столбцов
        for col in window.columns:
            window.tree.heading(col, text=col)
            window.tree.column(col, anchor="center", stretch=True, width=300)

        # Создание пролистывания по таблице
        scrollbar_y = ctk.CTkScrollbar(
            window.frame_table, width=15, height=50)
        scrollbar_y.place(x=985, y=0)
        window.tree.configure(yscrollcommand=scrollbar_y.set)

        # scrollbar_x = ctk.CTkScrollbar(window.frame_table, width=15, height=15)
        # scrollbar_x.place(x=0, y=295)
        # window.tree.configure(xscrollcommand=scrollbar_x.set)

        window.tree.place(x=0, y=0, relwidth=1, relheight=1)

        """--------------------------------------------------------ОШИБКА-ОТСУТСТВИЯ-ДАННЫХ-----------------------------------------------------------"""

        # Создание контейнера для ошибки при отсутствии вводных данных
        window.frame_false = ctk.CTkFrame(window, fg_color="white")

        # Отключение видимости ошибки
        window.frame_false.pack_forget()

        # Создание ярлыка отсутствия галочек
        ctk.CTkLabel(master=window.frame_false, text="Введите данные",
                      fg_color="white", text_color="red").pack(side="top")

        """--------------------------------------------------------ОШИБКА-ОТСУТСТВИЯ-СТУДЕНТОВ-----------------------------------------------------------"""

        # Создание контейнера для ошибки при отсутствии студентов
        window.frame_false_stud = ctk.CTkFrame(window, fg_color="white")

        # Отключение видимости ошибки
        window.frame_false_stud.pack_forget()

        # Создание ярлыка отсутствия галочек
        ctk.CTkLabel(master=window.frame_false_stud, text="Студент(ы) не найден(ы)",
                      fg_color="white", text_color="red").pack(side="top")

        """-------------------------------------------------------ОТКЛИК-НА-ТАБЛИЦУ----------------------------------------------------------"""

        # Возможность отклика на таблицу
        window.tree.bind("<Double-1>", window.click_on_table)

    # Реализация кнопки поика

    def click_find(window):

        if window.entry_familia.get() == "":
            if window.entry_name.get() == "":
                if window.entry_otchestvo.get() == "":
                    if window.day_var.get() == "1":
                        if window.month_var.get() == "1":
                            if window.year_var.get() == "1900":

                                window.frame_table.place_forget()
                                window.frame_false.place(x=440, y=290)
                                return

        window.frame_false_stud.place_forget()
        window.frame_false.place_forget()
        window.find()

    # Функция поиска
    def find(window):
        window.tree.delete(*window.tree.get_children())

        surname = window.entry_familia.get().upper() if window.entry_familia.get() else None
        name = window.entry_name.get().upper() if window.entry_name.get() else None
        och = window.entry_otchestvo.get().upper() if window.entry_otchestvo.get() else None
        day = window.day_var.get()
        month = window.month_var.get()
        year = window.year_var.get()
        birthdate = datetime(int(year), int(
            month), int(day))  # datetime(1990 01 02)
        global data
        global db

        data = db.get_person(surname, name, och, birthdate)

        if not data:
            window.frame_table.place_forget()
            window.frame_false_stud.place(x=440, y=290)
        else:
            for i, var in enumerate(data):
                window.tree.insert("", ctk.END, iid=f"I00{i}", values=var)
            window.frame_table.place(x=0, y=290)

    # Реализация изменения в таблице
    def click_on_table(window, event):
        row = window.tree.identify_row(event.y)
        global row_click
        row_click = int(row[1:])  # I000
        WindowInformation(window)


class WindowInformation(ctk.CTkToplevel):

    def __init__(window_inf, parent):
        super().__init__(parent)
        # Описание окна
        window_inf.title("Информация о студенте")
        center_window(window_inf, 900, 600)
        window_inf.configure(fg_color="white")
        window_inf.iconbitmap(ico_path)
        blue = "#6699CC"
        white = "white"
        black = "black"
        my_font = ctk.CTkFont(family="Arial", size=16)

        ctk.CTkFrame(window_inf, width=890, height=90,
                      fg_color=blue).place(x=5, y=5)
        window_inf.do_lable("Фамилия:", 10, 10, blue, blue, white)
        window_inf.do_lable(data[row_click][0], 80, 10, blue, blue, white) if data[row_click][0] else window_inf.do_lable(
            "Не указано", 80, 10, blue, blue, white)
        window_inf.do_lable("Имя:", 10, 40, blue, blue, white)
        window_inf.do_lable(data[row_click][2], 80, 40, blue, blue, white) if data[row_click][2] else window_inf.do_lable(
            "Не указано", 80, 40, blue, blue, white)
        window_inf.do_lable("Отчество:", 10, 70, blue, blue, white)
        window_inf.do_lable(data[row_click][4], 80, 70, blue, blue, white) if data[row_click][4] else window_inf.do_lable(
            "Не указано", 80, 70, blue, blue, white)

        window_inf.do_lable("Фамилия (лат):", 500, 10, blue, blue, white)
        window_inf.do_lable(data[row_click][1], 600, 10, blue, blue, white) if data[row_click][1] else window_inf.do_lable(
            "Не указано", 600, 10, blue, blue, white)
        window_inf.do_lable("Имя (лат):", 500, 40, blue, blue, white)
        window_inf.do_lable(data[row_click][3], 600, 40, blue, blue, white) if data[row_click][3] else window_inf.do_lable(
            "Не указано", 600, 40, blue, blue, white)
        window_inf.do_lable("Отчество (лат):", 500, 70, blue, blue, white)
        window_inf.do_lable(data[row_click][5], 600, 70, blue, blue, white) if data[row_click][5] else window_inf.do_lable(
            "Не указано", 600, 70, blue, blue, white)

        window_inf.do_lable("Пол:", 10, 100, white, white, black)
        window_inf.do_lable(data[row_click][8], 45, 100, white, white, black) if data[row_click][8] else window_inf.do_lable(
            "Не указано", 45, 100, white, white, black)
        window_inf.do_lable("Дата рождения:", 110, 100, white, white, black)
        window_inf.do_lable(data[row_click][7], 215, 100, white, white, black) if data[row_click][7] else window_inf.do_lable(
            "Не указано", 215, 100, white, white, black)
        window_inf.do_lable("Телефон:", 290, 100, white, white, black)
        window_inf.do_lable(f"+7{data[row_click][20]}", 350, 100, white, white,
                            black) if data[row_click][20] else window_inf.do_lable("Не указано", 350, 100, white, white, black)

        ctk.CTkFrame(window_inf, width=400, height=95,
                      fg_color=blue).place(x=5, y=125)
        window_inf.do_lable("Гражданство:", 10, 130, blue, blue, white)
        window_inf.do_lable(data[row_click][6], 100, 130, blue, blue, white) if data[row_click][6] else window_inf.do_lable(
            "Не указано", 100, 130, blue, blue, white)
        window_inf.do_lable("Государство рождения:",
                            10, 160, blue, blue, white)
        window_inf.do_lable(data[row_click][9], 160, 160, blue, blue, white) if data[row_click][9] else window_inf.do_lable(
            "Не указано", 160, 160, blue, blue, white)
        window_inf.do_lable("Город рождения:", 10, 190, blue, blue, white)
        window_inf.do_lable(data[row_click][10], 120, 190, blue, blue, white) if data[row_click][10] else window_inf.do_lable(
            "Не указано", 120, 190, blue, blue, white)

        ctk.CTkFrame(window_inf, width=400, height=90,
                      fg_color=blue).place(x=5, y=230)
        ctk.CTkLabel(master=window_inf, text="Паспортные данные", bg_color=blue,
                      fg_color=blue, text_color=white, font=my_font).place(x=10, y=235)

        window_inf.do_lable("Серия", 10, 265, blue, blue, white)
        window_inf.do_lable(data[row_click][11], 13, 285, blue, white, blue) if data[row_click][11] else window_inf.do_lable(
            "Не указано", 13, 285, blue, white, blue)
        window_inf.do_lable("Номер", 110, 265, blue, blue, white)
        window_inf.do_lable(data[row_click][12], 113, 285, blue, white, blue) if data[row_click][12] else window_inf.do_lable(
            "Не указано", 113, 285, blue, white, blue)
        window_inf.do_lable("Дата выдачи", 210, 265, blue, blue, white)
        window_inf.do_lable(data[row_click][13], 213, 285, blue, white, blue) if data[row_click][13] else window_inf.do_lable(
            "Не указано", 213, 285, blue, white, blue)
        window_inf.do_lable("Срок действия", 300, 265, blue, blue, white)
        window_inf.do_lable(data[row_click][14], 303, 285, blue, white, blue) if data[row_click][14] else window_inf.do_lable(
            "Не указано", 303, 285, blue, white, blue)

        if data[row_click][28] == "нет":
            ctk.CTkFrame(window_inf, width=400, height=90,
                          fg_color=blue).place(x=425, y=230)
            ctk.CTkLabel(master=window_inf, text="Виза", bg_color=blue,
                          fg_color=blue, text_color=white, font=my_font).place(x=430, y=235)
            window_inf.do_lable("Серия", 430, 265, blue, blue, white)
            window_inf.do_lable(data[row_click][16], 433, 285, blue, white, blue) if data[row_click][16] else window_inf.do_lable(
                "Не указано", 433, 285, blue, white, blue)
            window_inf.do_lable("Номер", 530, 265, blue, blue, white)
            window_inf.do_lable(data[row_click][17], 533, 285, blue, white, blue) if data[row_click][17] else window_inf.do_lable(
                "Не указано", 533, 285, blue, white, blue)
            window_inf.do_lable("Дата выдачи", 630, 265, blue, blue, white)
            window_inf.do_lable(data[row_click][18], 633, 285, blue, white, blue) if data[row_click][18] else window_inf.do_lable(
                "Не указано", 633, 285, blue, white, blue)
            window_inf.do_lable("Срок действия", 730, 265, blue, blue, white)
            window_inf.do_lable(data[row_click][19], 733, 285, blue, white, blue) if data[row_click][19] else window_inf.do_lable(
                "Не указано", 733, 285, blue, white, blue)
        elif data[row_click][28] == "ВНЖ":
            ctk.CTkFrame(window_inf, width=400, height=90,
                          fg_color=blue).place(x=425, y=230)
            ctk.CTkLabel(master=window_inf, text="ВНЖ", bg_color=blue,
                          fg_color=blue, text_color=white, font=my_font).place(x=430, y=235)
            window_inf.do_lable("Серия", 430, 265, blue, blue, white)
            window_inf.do_lable(data[row_click][31], 433, 285, blue, white, blue) if data[row_click][31] else window_inf.do_lable(
                "Не указано", 433, 285, blue, white, blue)
            window_inf.do_lable("Номер", 530, 265, blue, blue, white)
            window_inf.do_lable(data[row_click][32], 533, 285, blue, white, blue) if data[row_click][32] else window_inf.do_lable(
                "Не указано", 533, 285, blue, white, blue)
            window_inf.do_lable("Дата выдачи", 630, 265, blue, blue, white)
            window_inf.do_lable(data[row_click][29], 633, 285, blue, white, blue) if data[row_click][29] else window_inf.do_lable(
                "Не указано", 633, 285, blue, white, blue)
            window_inf.do_lable("Срок действия", 730, 265, blue, blue, white)
            window_inf.do_lable(data[row_click][30], 733, 285, blue, white, blue) if data[row_click][30] else window_inf.do_lable(
                "Не указано", 733, 285, blue, white, blue)
        elif data[row_click][28] == "РВПО":
            ctk.CTkFrame(window_inf, width=400, height=90,
                          fg_color=blue).place(x=425, y=230)
            ctk.CTkLabel(master=window_inf, text="РВПО", bg_color=blue,
                          fg_color=blue, text_color=white, font=my_font).place(x=430, y=235)
            window_inf.do_lable("Серия", 430, 265, blue, blue, white)
            window_inf.do_lable(data[row_click][31], 433, 285, blue, white, blue) if data[row_click][31] else window_inf.do_lable(
                "Не указано", 433, 285, blue, white, blue)
            window_inf.do_lable("Номер", 530, 265, blue, blue, white)
            window_inf.do_lable(data[row_click][32], 533, 285, blue, white, blue) if data[row_click][32] else window_inf.do_lable(
                "Не указано", 533, 285, blue, white, blue)
            window_inf.do_lable("Дата выдачи", 630, 265, blue, blue, white)
            window_inf.do_lable(data[row_click][29], 633, 285, blue, white, blue) if data[row_click][29] else window_inf.do_lable(
                "Не указано", 633, 285, blue, white, blue)
            window_inf.do_lable("Срок действия", 730, 265, blue, blue, white)
            window_inf.do_lable(data[row_click][30], 733, 285, blue, white, blue) if data[row_click][30] else window_inf.do_lable(
                "Не указано", 733, 285, blue, white, blue)

        window_inf.do_lable("Дата въезда", 10, 325, white, white, black)
        window_inf.do_lable(data[row_click][21], 13, 345, white, white, black) if data[row_click][21] else window_inf.do_lable(
            "Не указано", 13, 345, white, white, black)
        window_inf.do_lable("Срок пребывания до", 110,
                            325, white, white, black)
        window_inf.do_lable(data[row_click][22], 139, 345, white, white, black) if data[row_click][22] else window_inf.do_lable(
            "Не указано", 139, 345, white, white, black)

        ctk.CTkFrame(window_inf, width=400, height=90,
                      fg_color=blue).place(x=5, y=375)
        ctk.CTkLabel(master=window_inf, text="Миграционная карта", bg_color=blue,
                      fg_color=blue, text_color=white, font=my_font).place(x=10, y=380)
        window_inf.do_lable("Серия", 10, 410, blue, blue, white)
        window_inf.do_lable(data[row_click][23], 13, 430, blue, white, blue) if data[row_click][23] else window_inf.do_lable(
            "Не указано", 13, 430, blue, white, blue)
        window_inf.do_lable("Номер", 100, 410, blue, blue, white)
        window_inf.do_lable(data[row_click][24], 103, 430, blue, white, blue) if data[row_click][24] else window_inf.do_lable(
            "Не указано", 103, 430, blue, white, blue)

        ctk.CTkFrame(window_inf, width=400, height=90,
                      fg_color=blue).place(x=425, y=375)
        ctk.CTkLabel(master=window_inf, text="Общежитие", bg_color=blue,
                      fg_color=blue, text_color=white, font=my_font).place(x=430, y=380)
        num = data[row_click][25]
        window_inf.do_lable("Номер корпуса", 435, 410, blue, blue, white)
        window_inf.do_lable(num[0], 470, 430, blue, white, blue) if data[row_click][25] else window_inf.do_lable(
            "Не указано", 470, 430, blue, white, blue)
        window_inf.do_lable("Дата начала договора",
                            535, 410, blue, blue, white)
        window_inf.do_lable(data[row_click][26], 565, 430, blue, white, blue) if data[row_click][26] else window_inf.do_lable(
            "Не указано", 565, 430, blue, white, blue)
        window_inf.do_lable("Номер договора", 680, 410, blue, blue, white)
        window_inf.do_lable(data[row_click][27], 683, 430, blue, white, blue) if data[row_click][27] else window_inf.do_lable(
            "Не указано", 683, 430, blue, white, blue)

        # Создание кнопки "Сформировать"
        ctk.CTkButton(window_inf, text="Сформировать в Excel", text_color=white,
                       fg_color=blue, command=window_inf.click_form_excel).place(x=740, y=500)
        ctk.CTkButton(window_inf, text="Сформировать уведомление", text_color=white,
                       fg_color=blue, command=window_inf.click_form_Word).place(x=540, y=500)

        # Фиксация окна
        window_inf.grab_set()
        window_inf.focus_set()
        window_inf.wait_window()

    def do_lable(window_inf, text, x, y, bg_color, fg_color, text_color):
        ctk.CTkLabel(master=window_inf, text=text, height=10, fg_color=fg_color,
                      bg_color=bg_color, text_color=text_color).place(x=x, y=y)

    def click_form_excel(window_inf):
        WindowSaveExcel(window_inf)

    def click_form_Word(window_inf):
        WindowSaveWord(window_inf)


class WindowSaveExcel(ctk.CTkToplevel):
    def __init__(window_save, parent):
        super().__init__(parent)
        # Описание окна
        window_save.title("Сохранение в Exel")
        center_window(window_save, 300, 100)
        window_save.configure(fg_color="white")
        window_save.iconbitmap(ico_path)
        # Создание ярлыка сохранения
        ctk.CTkLabel(master=window_save, text="Сохранение", fg_color="white",
                      text_color="black", font=("Arial", 16)).pack(side="top", padx=7)

        # Создание контейнера для названия файла
        window_save.frame_name_file = ctk.CTkFrame(
            window_save, fg_color="white")
        window_save.frame_name_file.pack(fill=ctk.X)

        # Создание ярлыка и текстового поля для ввода названия файла
        ctk.CTkLabel(master=window_save.frame_name_file, text="Название файла:",
                      fg_color="white", text_color="black").grid(row=0, column=0, padx=7)
        window_save.entry_name_file = ctk.CTkEntry(
            master=window_save.frame_name_file, width=150)
        window_save.entry_name_file.insert(
            0, f"{data[row_click][1]} {data[row_click][3]}")

        window_save.entry_name_file.grid(row=0, column=1, pady=7)

        # Создание кнопки "Сохранить"
        window_save.search_button = ctk.CTkButton(
            window_save, text="Сохранить", command=window_save.click_save)
        window_save.search_button.pack(side="right", padx=5)

        # Обработка нажатия внутри области
        window_save.entry_name_file.bind("<FocusIn>", lambda e: window_save.entry_name_file.delete(
            0, len(window_save.entry_name_file.get())+1))

        # Фиксация окна
        window_save.grab_set()
        window_save.focus_set()
        window_save.wait_window()

    def click_save(window_save):
        CompletionExcel(window_save.entry_name_file.get())
        window_save.destroy()


class WindowSaveWord(ctk.CTkToplevel):
    def __init__(window_save, parent):
        super().__init__(parent)
        # Описание окна
        window_save.title("Сохранение в Exel")
        center_window(window_save, 500, 250)
        window_save.configure(fg_color="white")
        window_save.iconbitmap(ico_path)
        # Создание ярлыка сохранения
        ctk.CTkLabel(master=window_save, text="Сохранение", fg_color="white",
                      text_color="white", font=("Arial", 16)).place(x=210, y=10)

        window_save.var_check_1 = ctk.BooleanVar()
        window_save.var_check_2 = ctk.BooleanVar()
        window_save.var_check_3 = ctk.BooleanVar()

        # # Создание галочек
        ctk.CTkCheckBox(window_save, variable=window_save.var_check_1,
                         bg_color="white").place(x=10, y=50)
        ctk.CTkLabel(master=window_save, text="о предоставлении иностранному гражданину (лицу без гражданства)\n"
                      "академического отпуска образовательной/научной организации;", bg_color="white").place(x=40, y=50)
        ctk.CTkCheckBox(window_save, variable=window_save.var_check_2,
                         bg_color="white").place(x=10, y=90)
        ctk.CTkLabel(master=window_save, text="о досрочном прекращении обучения иностранного гражданина\n"
                      "(лица без гражданства) в образовательной/научной организации;", bg_color="white").place(x=40, y=90)
        ctk.CTkCheckBox(window_save, variable=window_save.var_check_3,
                         bg_color="white").place(x=10, y=130)
        ctk.CTkLabel(master=window_save, text="о завершении обучения иностранного гражданина (лица без гражданства)\n"
                      "(лица без гражданства) в образовательной/научной организации;", bg_color="white").place(x=40, y=130)

        # Создание ярлыка и текстового поля для ввода названия файла
        ctk.CTkLabel(master=window_save, text="Название файла:",
                      bg_color="white").place(x=30, y=180)
        window_save.entry_name_file = ctk.CTkEntry(
            master=window_save, width=300)
        window_save.entry_name_file.insert(
            0, f"Уведомление {data[row_click][1]} {data[row_click][3]}")
        window_save.entry_name_file.place(x=30, y=200)

        # Создание кнопки "Сохранить"
        window_save.search_button = ctk.CTkButton(
            window_save, text="Сохранить", command=window_save.click_save)
        window_save.search_button.place(x=350, y=199)

        # Обработка нажатия внутри области
        window_save.entry_name_file.bind("<FocusIn>", lambda e: window_save.entry_name_file.delete(
            0, len(window_save.entry_name_file.get())+1))

        # Фиксация окна
        window_save.grab_set()
        window_save.focus_set()
        window_save.wait_window()

    def click_save(window_save):
        ComplectionWord(window_save.entry_name_file.get(), window_save)
        window_save.destroy()


# Функция центрирования окна
def center_window(
                self, 
                width: int, 
                height: int
) -> None:

    # Получаем размеры экрана
    screen_width = self.winfo_screenwidth()
    screen_height = self.winfo_screenheight()

    # Вычисляем координаты верхнего левого угла
    x = (screen_width - width) // 2
    y = (screen_height - height) // 2

    # Устанавливаем размеры и положение окна
    self.geometry(f"{width}x{height}+{x}+{y}")


# Функция Заполнения Excel-таблицы
def change_sheet(
                sheet, 
                row: int, 
                col: int, 
                step: int, 
                value: str | None, 
                max_col: int
) -> None:
    col_cur = col
    row_cur = row
    for val in value:
        if col_cur <= max_col:
            sheet.cell(row=row_cur, column=col_cur, value=val)
            col_cur += step


def CompletionExcel(file_out: str) -> None:
    # Загрузка книги
    wb = openpyxl.load_workbook(excle_path)
    # Работаем с активным листом (если нужен другой, уточните)
    sheet = wb.active

    # Начало заполнения таблицы
    # Фамилия на русском
    change_sheet(sheet, 10, 14, 4, data[row_click][0],  122)
    change_sheet(sheet, 176, 14, 4, data[row_click][0],  122)
    # Фамилия на латинице
    change_sheet(sheet, 12, 38, 4, data[row_click][1],  122)
    # Имя на русском
    change_sheet(sheet, 14, 14, 4, data[row_click][2],  122)
    change_sheet(sheet, 178, 14, 4, data[row_click][2],  122)
    # Имя на латинице
    change_sheet(sheet, 16, 34, 4, data[row_click][3],  122)
    # Отчество на русском
    change_sheet(sheet, 18, 30, 4, data[row_click][4],  122)
    change_sheet(sheet, 180, 34, 4, data[row_click][4],  122)
    # Отчество на латинице
    change_sheet(sheet, 20, 38, 4, data[row_click][5],  122)
    # Гражданство
    change_sheet(sheet, 24, 22, 4, data[row_click][6],  122)
    change_sheet(sheet, 182, 18, 4, data[row_click][6],  122)
    # Дата рождения
    date = data[row_click][7]
    # Число рождения
    change_sheet(sheet, 27, 30, 0, date[0],  30)
    change_sheet(sheet, 27, 34, 0, date[1],  34)
    change_sheet(sheet, 185, 30, 0, date[0],  30)
    change_sheet(sheet, 185, 34, 0, date[1],  34)
    # Месяц рождения
    change_sheet(sheet, 27, 46, 0, date[3],  46)
    change_sheet(sheet, 27, 50, 0, date[4],  50)
    change_sheet(sheet, 185, 50, 0, date[3],  50)
    change_sheet(sheet, 185, 54, 0, date[4],  54)
    # Год рождения
    change_sheet(sheet, 27, 58, 0, date[6],  58)
    change_sheet(sheet, 27, 62, 0, date[7],  62)
    change_sheet(sheet, 27, 66, 0, date[8],  66)
    change_sheet(sheet, 27, 70, 0, date[9],  70)
    change_sheet(sheet, 185, 66, 0, date[6],  66)
    change_sheet(sheet, 185, 70, 0, date[7],  70)
    change_sheet(sheet, 185, 74, 0, date[8],  74)
    change_sheet(sheet, 185, 78, 0, date[9],  78)
    # Пол
    if data[row_click][8] == "Мужской":
        change_sheet(sheet, 27, 90, 0, "X",  90)
        change_sheet(sheet, 185, 102, 0, "X",  102)
    elif data[row_click][8] == "Женский":
        change_sheet(sheet, 27, 106, 0, "X",  106)
        change_sheet(sheet, 185, 118, 0, "X",  118)
    # Государство рождения
    change_sheet(sheet, 29, 26, 4, data[row_click][9],  122)
    # Город рождения
    change_sheet(sheet, 31, 2, 4, data[row_click][10], 122)
    # Серия паспорта
    change_sheet(sheet, 37, 10, 4, data[row_click][11],  30)
    change_sheet(sheet, 187, 58, 4, data[row_click][11],  78)
    # Номер паспорта
    change_sheet(sheet, 37, 42, 4, data[row_click][12],  122)
    change_sheet(sheet, 187, 86, 4, data[row_click][12],  122)
    # Дата выдачи паспорта
    date = data[row_click][13]
    # Число
    change_sheet(sheet, 39, 9, 0, date[0],  9)
    change_sheet(sheet, 39, 13, 0, date[1],  13)
    change_sheet(sheet, 189, 9, 0, date[0],  9)
    change_sheet(sheet, 189, 13, 0, date[1],  13)
    # Месяц выдачи
    change_sheet(sheet, 39, 26, 0, date[3],  26)
    change_sheet(sheet, 39, 30, 0, date[4],  30)
    change_sheet(sheet, 189, 26, 0, date[3],  26)
    change_sheet(sheet, 189, 30, 0, date[4],  30)
    # Год выдачи
    change_sheet(sheet, 39, 38, 0, date[6],  38)
    change_sheet(sheet, 39, 42, 0, date[7],  42)
    change_sheet(sheet, 39, 46, 0, date[8],  46)
    change_sheet(sheet, 39, 50, 0, date[9],  50)
    change_sheet(sheet, 189, 38, 0, date[6],  38)
    change_sheet(sheet, 189, 42, 0, date[7],  42)
    change_sheet(sheet, 189, 46, 0, date[8],  46)
    change_sheet(sheet, 189, 50, 0, date[9],  50)
    # Дата срока действия паспорта
    date = data[row_click][14]
    # Число
    change_sheet(sheet, 39, 66, 0, date[0],  66)
    change_sheet(sheet, 39, 70, 0, date[1],  70)
    change_sheet(sheet, 189, 66, 0, date[0],  66)
    change_sheet(sheet, 189, 70, 0, date[1],  70)
    # Месяц
    change_sheet(sheet, 39, 82, 0, date[3],  82)
    change_sheet(sheet, 39, 86, 0, date[4],  86)
    change_sheet(sheet, 189, 82, 0, date[3],  82)
    change_sheet(sheet, 189, 86, 0, date[4],  86)
    # Год
    change_sheet(sheet, 39, 94, 0, date[6],  94)
    change_sheet(sheet, 39, 98, 0, date[7],  98)
    change_sheet(sheet, 39, 102, 0, date[8],  102)
    change_sheet(sheet, 39, 106, 0, date[9],  106)
    change_sheet(sheet, 189, 94, 0, date[6],  94)
    change_sheet(sheet, 189, 98, 0, date[7],  98)
    change_sheet(sheet, 189, 102, 0, date[8],  102)
    change_sheet(sheet, 189, 106, 0, date[9],  106)
    # Проверка на наличие визы, ВНЖ, РВПО
    if data[row_click][28] == "нет":
        # Наличие визы
        if data[row_click][15] == "1":
            change_sheet(sheet, 53, 8, 0, "X",  8)
        # Серия визы
        change_sheet(sheet, 56, 10, 4, data[row_click][16],  30)
        # Номер визы
        change_sheet(sheet, 56, 42, 4, data[row_click][17],  122)
        # Дата выдачи визы
        date = data[row_click][18]
        # Число
        change_sheet(sheet, 58, 9, 0, date[0],  9)
        change_sheet(sheet, 58, 13, 0, date[1],  13)
        # Месяц
        change_sheet(sheet, 58, 26, 0, date[3],  26)
        change_sheet(sheet, 58, 30, 0, date[4],  30)
        # Год
        change_sheet(sheet, 58, 38, 0, date[6],  38)
        change_sheet(sheet, 58, 42, 0, date[7],  42)
        change_sheet(sheet, 58, 46, 0, date[8],  46)
        change_sheet(sheet, 58, 50, 0, date[9],  50)
        # Дата срока визы
        date = data[row_click][19]
        # Число
        change_sheet(sheet, 58, 66, 0, date[0],  66)
        change_sheet(sheet, 58, 70, 0, date[1],  70)
        # Месяц
        change_sheet(sheet, 58, 82, 0, date[3],  82)
        change_sheet(sheet, 58, 86, 0, date[4],  86)
        # Год
        change_sheet(sheet, 58, 94, 0, date[6],  94)
        change_sheet(sheet, 58, 98, 0, date[7],  98)
        change_sheet(sheet, 58, 102, 0, date[8],  102)
        change_sheet(sheet, 58, 106, 0, date[9],  106)
    elif data[row_click][28] == "РВПО":
        change_sheet(sheet, 53, 108, 0, "X",  108)
        # Серия РВПО
        change_sheet(sheet, 56, 10, 4, data[row_click][31],  30)
        # Номер РВПО
        change_sheet(sheet, 56, 42, 4, data[row_click][32],  122)
        # Дата выдачи РВПО
        date = data[row_click][29]
        # Число
        change_sheet(sheet, 58, 9, 0, date[0],  9)
        change_sheet(sheet, 58, 13, 0, date[1],  13)
        # Месяц
        change_sheet(sheet, 58, 26, 0, date[3],  26)
        change_sheet(sheet, 58, 30, 0, date[4],  30)
        # Год
        change_sheet(sheet, 58, 38, 0, date[6],  38)
        change_sheet(sheet, 58, 42, 0, date[7],  42)
        change_sheet(sheet, 58, 46, 0, date[8],  46)
        change_sheet(sheet, 58, 50, 0, date[9],  50)
        # Дата срока РВПО
        date = data[row_click][30]
        # Число
        change_sheet(sheet, 58, 66, 0, date[0],  66)
        change_sheet(sheet, 58, 70, 0, date[1],  70)
        # Месяц
        change_sheet(sheet, 58, 82, 0, date[3],  82)
        change_sheet(sheet, 58, 86, 0, date[4],  86)
        # Год
        change_sheet(sheet, 58, 94, 0, date[6],  94)
        change_sheet(sheet, 58, 98, 0, date[7],  98)
        change_sheet(sheet, 58, 102, 0, date[8],  102)
        change_sheet(sheet, 58, 106, 0, date[9],  106)
    elif data[row_click][28] == "ВНЖ":
        change_sheet(sheet, 53, 36, 0, "X",  36)
        # Серия ВНЖ
        change_sheet(sheet, 56, 10, 4, data[row_click][31],  30)
        # Номер ВНЖ
        change_sheet(sheet, 56, 42, 4, data[row_click][32],  122)
        # Дата выдачи ВНЖ
        date = data[row_click][29]
        # Число
        change_sheet(sheet, 58, 9, 0, date[0],  9)
        change_sheet(sheet, 58, 13, 0, date[1],  13)
        # Месяц
        change_sheet(sheet, 58, 26, 0, date[3],  26)
        change_sheet(sheet, 58, 30, 0, date[4],  30)
        # Год
        change_sheet(sheet, 58, 38, 0, date[6],  38)
        change_sheet(sheet, 58, 42, 0, date[7],  42)
        change_sheet(sheet, 58, 46, 0, date[8],  46)
        change_sheet(sheet, 58, 50, 0, date[9],  50)
        # Дата срока ВНЖ
        date = data[row_click][30]
        # Число
        change_sheet(sheet, 58, 66, 0, date[0],  66)
        change_sheet(sheet, 58, 70, 0, date[1],  70)
        # Месяц
        change_sheet(sheet, 58, 82, 0, date[3],  82)
        change_sheet(sheet, 58, 86, 0, date[4],  86)
        # Год
        change_sheet(sheet, 58, 94, 0, date[6],  94)
        change_sheet(sheet, 58, 98, 0, date[7],  98)
        change_sheet(sheet, 58, 102, 0, date[8],  102)
        change_sheet(sheet, 58, 106, 0, date[9],  106)
    # Телефон
    # change_sheet(sheet, 62, 82, 4, data[row_click][20],  118)
    # Дата въезда
    date = data[row_click][21]
    # Число
    change_sheet(sheet, 66, 9, 0, date[0],  9)
    change_sheet(sheet, 66, 13, 0, date[1],  13)
    # Месяц
    change_sheet(sheet, 66, 26, 0, date[3],  26)
    change_sheet(sheet, 66, 30, 0, date[4],  30)
    # Год
    change_sheet(sheet, 66, 38, 0, date[6],  38)
    change_sheet(sheet, 66, 42, 0, date[7],  42)
    change_sheet(sheet, 66, 46, 0, date[8],  46)
    change_sheet(sheet, 66, 50, 0, date[9],  50)
    # Дата срок пребывания
    date = data[row_click][22]
    # Число
    change_sheet(sheet, 66, 66, 0, date[0],  66)
    change_sheet(sheet, 66, 70, 0, date[1],  70)
    change_sheet(sheet, 222, 49, 0, date[0],  49)
    change_sheet(sheet, 222, 53, 0, date[1],  53)
    change_sheet(sheet, 266, 38, 0, date[0],  38)
    change_sheet(sheet, 266, 42, 0, date[1],  42)
    # Месяц
    change_sheet(sheet, 66, 82, 0, date[3],  82)
    change_sheet(sheet, 66, 86, 0, date[4],  86)
    change_sheet(sheet, 222, 65, 0, date[3],  65)
    change_sheet(sheet, 222, 69, 0, date[4],  69)
    change_sheet(sheet, 266, 54, 0, date[3],  54)
    change_sheet(sheet, 266, 58, 0, date[4],  58)
    # Год
    change_sheet(sheet, 66, 94, 0, date[6],  94)
    change_sheet(sheet, 66, 98, 0, date[7],  98)
    change_sheet(sheet, 66, 102, 0, date[8],  102)
    change_sheet(sheet, 66, 106, 0, date[9],  106)
    change_sheet(sheet, 222, 79, 0, date[6],  79)
    change_sheet(sheet, 222, 83, 0, date[7],  83)
    change_sheet(sheet, 222, 87, 0, date[8],  87)
    change_sheet(sheet, 222, 91, 0, date[9],  91)
    change_sheet(sheet, 266, 68, 0, date[6],  68)
    change_sheet(sheet, 266, 72, 0, date[7],  72)
    change_sheet(sheet, 266, 76, 0, date[8],  76)
    change_sheet(sheet, 266, 80, 0, date[9],  80)
    # Миграционная карта серия
    change_sheet(sheet, 68, 34, 4, data[row_click][23],  46)
    # Миграционная карта номер
    change_sheet(sheet, 68, 54, 4, data[row_click][24],  102)

    # Номер корпус
    if data[row_click][25]:
        num = data[row_click][25]
        sheet.cell(row=104, column=46, value=f"КОРПУС {num[0]}")
        sheet.cell(row=200, column=46, value=f"КОРПУС {num[0]}")
    # Дата начала найма
    date_naym = data[row_click][26]
    if date_naym != "":
        # Число найма
        change_sheet(sheet, 112, 90, 0, date_naym[0],  90)
        change_sheet(sheet, 112, 94, 0, date_naym[1],  94)
        # Месяц найма
        change_sheet(sheet, 112, 102, 0, date_naym[3],  102)
        change_sheet(sheet, 112, 106, 0, date_naym[4],  106)
        # Год найма
        change_sheet(sheet, 112, 114, 0, date_naym[8],  114)
        change_sheet(sheet, 112, 118, 0, date_naym[9],  118)
        # Номер договора
        change_sheet(sheet, 114, 46, 4, data[row_click][27],  90)
    if data[row_click][42]:
        pattern_reg = r"РЕСПУБЛИКА\s[А-ЯЁ]+(?:\s[А-ЯЁ]+)?|[А-ЯЁ]+(?:-[А-ЯЁ]+)?\s(?:ОБЛАСТЬ|КРАЙ|РЕСПУБЛИКА|АО)"
        pattern_rayon = r"[А-ЯЁ]+(?:-[А-ЯЁ]+)?\s(?:РАЙОН)"
        pattern_city = r"(?:Г|ПГТ|РП|КП|К|ДП|П|НП|С|М|Д|СЛ|СТ|СТ-ЦА|Х|КЛХ|СВХ)\.\s?\b[А-ЯЁ]+(?:-[А-ЯЁ]+)?\b"
        pattern_street = r"(?:УЛ|АЛЛЕЯ|БУЛЬВАР|МАГИСТРАЛЬ|ПЕРЕУЛОК|ПЛОЩАДЬ|ПРОЕЗД|ПРОСПЕКТ|ПРОУЛОК|РАЗЪЕЗД|СПУСК|ТРАКТ|ТУПИК|ШОССЕ|УЛИЦА)\.?\s?[А-ЯЁ]+(?:-[А-ЯЁ]+)?|[А-ЯЁ]+(?:-[А-ЯЁ]+)?\s(?:УЛ|ПР|П|ПРОЕЗД|УЛИЦА|ПРОСПЕКТ|БУЛЬВАР)\.?"
        pattern_dom = r"Д.\s?\b[0-9]+(?:[А-ЯЁ]+)?\b"
        pattern_korp = r"КОРП.\s?\b[0-9]+\b|КОРП.\s?\b[А-ЯЁ]+\b"
        pattern_str = r"СТР.\s?\b[0-9]+\b|СТР.\s?\b[А-ЯЁ]+\b"
        pattern_kv = r"КВ.\s?\b[0-9]+\b"
        reg = "".join(re.findall(pattern_reg, data[row_click][42]))
        change_sheet(sheet, 83, 2, 4, reg, 122)
        rayon = "".join(re.findall(pattern_rayon, data[row_click][42]))
        change_sheet(sheet, 85, 2, 4, rayon, 122)
        city = "".join(re.findall(pattern_city, data[row_click][42]))
        change_sheet(sheet, 87, 2, 4, city, 122)
        street = "".join(re.findall(pattern_street, data[row_click][42]))
        change_sheet(sheet, 89, 2, 4, street, 122)
        dom = "".join(re.findall(pattern_dom, data[row_click][42]))
        sheet.cell(row=91, column=2, value=f"ДОМ {dom.replace('Д. ','')}")
        korp = "".join(re.findall(pattern_korp, data[row_click][42]))
        str_ = "".join(re.findall(pattern_str, data[row_click][42]))
        if korp:
            sheet.cell(row=91, column=46,
                       value=f"КОРПУС {korp.replace('КОРП. ','')}")
        elif str_:
            sheet.cell(row=91, column=46,
                       value=f"СТРОЕНИЕ {str_.replace('СТР. ','')}")
        kv = "".join(re.findall(pattern_kv, data[row_click][42]))
        if kv:

            sheet.cell(row=93, column=2,
                       value=f"КВАРТИРА {kv.replace('КВ. ','')}")

    # Сохранение изменений
    if os.path.isdir(".\\out"):
        pass
    else:
        os.mkdir(f".\\out")

    output_path = f".\\out\\{file_out}.xlsx"
    wb.save(output_path)
    os.startfile(output_path)


def ComplectionWord(file_out: str, window_save: ctk.CTkToplevel) -> None:
    print(doc_path)
    doc = Document(doc_path)
    if window_save.var_check_1.get() == True:
        UpdateWord(doc, 1, 12, 1, "X")
    if window_save.var_check_2.get() == True:
        UpdateWord(doc, 1, 14, 1, "X")
    if window_save.var_check_3.get() == True:
        UpdateWord(doc, 1, 16, 1, "X")
    UpdateWord(doc, 1, 19, 3, data[row_click][0])  # Фамилия на русском
    UpdateWord(doc, 1, 19, 15, data[row_click][1])  # Фамилия на латинице
    UpdateWord(doc, 1, 21, 4, data[row_click][2])  # Имя на русском
    UpdateWord(doc, 1, 21, 15, data[row_click][3])  # Имя на латинице
    UpdateWord(doc, 1, 23, 7, data[row_click][4])  # Отчество на русском
    UpdateWord(doc, 1, 23, 15, data[row_click][5])  # Отчество на латинице
    UpdateWord(doc, 1, 25, 5, data[row_click][7])  # Дата рождения
    # Стаана и город рождения
    UpdateWord(doc, 1, 25, 13,
               f"{data[row_click][9]}, {data[row_click][10]}", 11)
    UpdateWord(doc, 1, 27, 8, data[row_click][6])  # Гражданство
    if data[row_click][8] == "Мужской":
        UpdateWord(doc, 1, 27, 20, "X")  # Пол
    elif data[row_click][8] == "Женский":
        UpdateWord(doc, 1, 27, 22, "X")  # Пол
    UpdateWord(doc, 1, 29, 2, data[row_click][11])  # Серия паспорта
    UpdateWord(doc, 1, 29, 6, data[row_click][12])  # Номер паспорта
    UpdateWord(doc, 1, 29, 15, data[row_click][13], 11)  # Дата выдачи паспорта
    # Дата срока действия паспорта
    UpdateWord(doc, 1, 29, 19, data[row_click][14], 11)
    if data[row_click][25]:
        num = data[row_click][25]
        UpdateWord(
            doc, 2, 2, 0, f"Г. МОСКВА, КОЧНОВСКИЙ ПР., Д.7, КОРПУС {num[0]}")
    UpdateWord(doc, 2, 3, 14, data[row_click][21])  # Дата въезда
    UpdateWord(doc, 2, 3, 23, data[row_click][22])  # Срок пребывания
    if data[row_click][28] == "нет":
        UpdateWord(doc, 2, 6, 4, data[row_click]
                   [33], 12, False)  # Кратност визы
        UpdateWord(doc, 2, 6, 28, "учебная", 12, False)  # Категория визы
        UpdateWord(doc, 2, 7, 3, "учеба", 12, False)  # Цель визы
        UpdateWord(doc, 2, 7, 19, data[row_click][16])  # Серия визы
        UpdateWord(doc, 2, 7, 28, data[row_click][17])  # Номер визы
        # Инденитификационный номер визы
        UpdateWord(doc, 2, 8, 8, data[row_click][34])
        UpdateWord(doc, 2, 8, 24, data[row_click][18])  # Дата выдачи визы
        UpdateWord(doc, 2, 8, 31, data[row_click][19])  # Дата срока визы
    elif data[row_click][28] == "РВПО":
        if data[row_click][16] != "" or data[row_click][16].isspace():
            UpdateWord(doc, 2, 6, 4, data[row_click]
                       [33], 12, False)  # Кратност визы
            UpdateWord(doc, 2, 6, 28, "учебная", 12, False)  # Категория визы
            UpdateWord(doc, 2, 7, 3, "учеба", 12, False)  # Цель визы
            UpdateWord(doc, 2, 7, 19, data[row_click][16])  # Серия визы
            UpdateWord(doc, 2, 7, 28, data[row_click][17])  # Номер визы
            # Инденитификационный номер визы
            UpdateWord(doc, 2, 8, 8, data[row_click][34])
            UpdateWord(doc, 2, 8, 24, data[row_click][18])  # Дата выдачи визы
            UpdateWord(doc, 2, 8, 31, data[row_click][19])  # Дата срока визы
    elif data[row_click][28] == "ВНЖ":
        if data[row_click][16] != "" or data[row_click][16].isspace():
            UpdateWord(doc, 2, 6, 4, data[row_click]
                       [33], 12, False)  # Кратност визы
            UpdateWord(doc, 2, 6, 28, "учебная", 12, False)  # Категория визы
            UpdateWord(doc, 2, 7, 3, "учеба", 12, False)  # Цель визы
            UpdateWord(doc, 2, 7, 19, data[row_click][16])  # Серия визы
            UpdateWord(doc, 2, 7, 28, data[row_click][17])  # Номер визы
            # Инденитификационный номер визы
            UpdateWord(doc, 2, 8, 8, data[row_click][34])
            UpdateWord(doc, 2, 8, 24, data[row_click][18])  # Дата выдачи визы
            UpdateWord(doc, 2, 8, 31, data[row_click][19])  # Дата срока визы
    if data[row_click][38] == "" or data[row_click][38].isspace():
        year_kon_st = 0
    else:
        year_kon_st = int(data[row_click][38][6]+data[row_click]
                          [38][7]+data[row_click][38][8]+data[row_click][38][9])
    if data[row_click][40] == "" or data[row_click][40].isspace():
        year_gos_st = 0
    else:
        year_gos_st = int(data[row_click][40][6]+data[row_click]
                          [40][7]+data[row_click][40][8]+data[row_click][40][9])
    if year_gos_st > year_kon_st:
        UpdateWord(doc, 2, 12, 10, "гос.направление", 11, False)  # Направление
        # Дата выдачи контракта
        UpdateWord(doc, 2, 12, 21, data[row_click][40], 11)
        UpdateWord(doc, 2, 12, 28, data[row_click][35], 11)  # Номер контракта
        UpdateWord(doc, 2, 14, 5, data[row_click][40], 11)  # Срок обучения с
        UpdateWord(doc, 2, 14, 15, data[row_click][41], 11)  # Срок обучения по
    else:
        UpdateWord(doc, 2, 12, 10, "контракт", 11, False)  # Направление
        # Дата выдачи контракта
        UpdateWord(doc, 2, 12, 21, data[row_click][38], 11)
        UpdateWord(doc, 2, 12, 28, data[row_click][37], 11)  # Номер контракта
        UpdateWord(doc, 2, 14, 5, data[row_click][38], 11)  # Срок обучения с
        UpdateWord(doc, 2, 14, 15, data[row_click][39], 11)  # Срок обучения по

    # Сохранение изменений
    if os.path.isdir(".\\out"):
        pass
    else:
        os.mkdir(f".\\out")

    output_path = f".\\out\\{file_out}.docx"
    doc.save(output_path)
    os.startfile(output_path)


def UpdateWord(
                doc, 
                table_index: int, 
                row_index: int, 
                cell_index: int, 
                new_text: str, 
                font_size = 12 , 
                uppercase = True
) -> None:
    """ Обновляет текст в ячейке таблицы, добавляя жирность и изменение регистра """
    try:
        table = doc.tables[table_index]
        if row_index < len(table.rows) and cell_index < len(table.rows[row_index].cells):
            cell = table.rows[row_index].cells[cell_index]

            # Определяем окончательный текст
            formatted_text = new_text.upper() if uppercase else new_text.lower()

            # Если в ячейке нет параграфов, создаем новый
            if not cell.paragraphs:
                paragraph = cell.add_paragraph()
            else:
                paragraph = cell.paragraphs[0]

            # Если нет runs, создаем новый run
            if not paragraph.runs:
                run = paragraph.add_run(formatted_text)
            else:
                run = paragraph.runs[0]
                run.text = formatted_text

            # Применяем жирный стиль
            run.bold = True
            run.font.name = "Times New Roman"  # Устанавливаем шрифт
            run.font.size = Pt(font_size)  # Устанавливаем размер шрифта
        else:
            print(f"Invalid row or cell index: ({row_index}, {cell_index})")
    except IndexError as e:
        print(f"Error accessing table: {e}")


def start_app():
    app = WindowProgramm()
    app.withdraw()
    UpdateApp(__version__)
    WindowAuthorizationConnectionBD(app)
    app.mainloop()
    del app


# Начало работы программы
if __name__ == "__main__":

    start_app()
